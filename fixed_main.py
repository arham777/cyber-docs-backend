import os
import shutil
from docx import Document
from docx.shared import Inches  # Added import for Inches
import fitz  # PyMuPDF
from zipfile import ZipFile
import tempfile
import xml.etree.ElementTree as ET
import re
from pdf2docx import Converter  # Add this import

# Configuration
TEMPLATE_DOCX = "cybergen-template.docx"  # Your branded template
INPUT_DIR = "input"
OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Register XML namespaces to prevent prefix generation
namespaces = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "p": "http://schemas.openxmlformats.org/package/2006/relationships",
    "ct": "http://schemas.openxmlformats.org/package/2006/content-types"
}

# Register namespaces for proper XML generation
for prefix, uri in namespaces.items():
    ET.register_namespace(prefix, uri)

def update_section_margins(sect_pr_element):
    """Update or add margin settings to a section properties element."""
    # Define standard margins (in twentieths of a point)
    margins = {
        "top": "1440",      # 1 inch = 1440 twentieths of a point
        "right": "1440",    # 1 inch right margin
        "bottom": "1440",   # 1 inch bottom margin
        "left": "1440",     # 1 inch left margin
        "header": "720",    # 0.5 inch for header
        "footer": "720"     # 0.5 inch for footer
    }
    
    # Create new pgMar element with proper namespace
    pg_mar = ET.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgMar")
    
    # Set all margin values
    for margin_type, value in margins.items():
        pg_mar.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}" + margin_type, value)
    
    # Create pgSz with A4 dimensions
    pg_sz = ET.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgSz")
    pg_sz.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w", "11906")  # A4 width (210mm)
    pg_sz.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}h", "16838")  # A4 height (297mm)
    pg_sz.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}orient", "portrait")
    
    # First remove all existing elements to ensure proper order
    for child in list(sect_pr_element):
        if child.tag.endswith('}pgSz') or child.tag.endswith('}pgMar'):
            sect_pr_element.remove(child)
    
    # Add elements in correct order
    has_type = False
    for child in sect_pr_element:
        if child.tag.endswith('}type'):
            has_type = True
            break
    
    if not has_type:
        type_element = ET.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type")
        type_element.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "nextPage")
        sect_pr_element.insert(0, type_element)
    
    # Insert in correct order (after type element)
    # Note: Proper order is important for Word to read the document correctly
    sect_pr_element.insert(1, pg_sz)
    sect_pr_element.insert(2, pg_mar)
    
    return sect_pr_element

def create_consistent_document(input_path, output_path):
    """Apply template's headers/footers to a DOCX file using a simpler approach."""
    try:
        # Create a copy of the template
        shutil.copy2(TEMPLATE_DOCX, output_path)
        
        # Open both documents
        template_doc = Document(output_path)  # This is now a copy of the template
        content_doc = Document(input_path)    # The document with content
        
        # Clear all paragraphs from the template while keeping headers/footers
        for section in template_doc.sections:
            section.different_first_page_header_footer = False
        
        # Clear all content from the template copy
        for paragraph in template_doc.paragraphs:
            p = paragraph._p
            p.getparent().remove(p)
        
        for table in template_doc.tables:
            tbl = table._tbl
            tbl.getparent().remove(tbl)
        
        # Add all content from the content document to the template
        for paragraph in content_doc.paragraphs:
            p = template_doc.add_paragraph()
            # Copy paragraph text and formatting
            p.text = paragraph.text
            p.style = paragraph.style
            
            # Copy runs with formatting
            if len(paragraph.runs) > 0:
                p.text = ""  # Clear text to avoid duplication
                for run in paragraph.runs:
                    new_run = p.add_run(run.text)
                    # Copy formatting
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    if run.font.size:
                        new_run.font.size = run.font.size
                    if run.font.color.rgb:
                        new_run.font.color.rgb = run.font.color.rgb
        
        # Add all tables
        for table in content_doc.tables:
            new_table = template_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            # Copy table contents
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                        new_table.rows[i].cells[j].text = cell.text
        
        # Save the combined document
        template_doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error in simplified approach: {e}")
        return False

def apply_branding_to_docx(input_path, output_path):
    """Apply template's headers/footers to a DOCX file using low-level zip manipulation."""
    # First try using the simpler approach
    if create_consistent_document(input_path, output_path):
        print("Successfully created document using simplified approach")
        return
    
    print("Falling back to advanced approach...")
    
    # Create temporary directories
    with tempfile.TemporaryDirectory() as temp_dir:
        template_dir = os.path.join(temp_dir, "template")
        target_dir = os.path.join(temp_dir, "target")
        
        os.makedirs(template_dir, exist_ok=True)
        os.makedirs(target_dir, exist_ok=True)
        
        # Extract both docx files (they're just zip files)
        with ZipFile(TEMPLATE_DOCX, 'r') as template_zip:
            template_zip.extractall(template_dir)
            
        with ZipFile(input_path, 'r') as target_zip:
            target_zip.extractall(target_dir)
        
        # Copy document.xml.rels first (this is critical for proper relationship IDs)
        template_doc_rels_file = os.path.join(template_dir, "word", "_rels", "document.xml.rels")
        target_doc_rels_file = os.path.join(target_dir, "word", "_rels", "document.xml.rels")
        
        if os.path.exists(template_doc_rels_file):
            # Ensure target directory exists
            os.makedirs(os.path.dirname(target_doc_rels_file), exist_ok=True)
            shutil.copy2(template_doc_rels_file, target_doc_rels_file)
            print("Copied document relationships from template")
        
        # Copy all header and footer files from template to target
        template_word_dir = os.path.join(template_dir, "word")
        target_word_dir = os.path.join(target_dir, "word")
        
        for filename in os.listdir(template_word_dir):
            if filename.startswith(("header", "footer")):
                source_file = os.path.join(template_word_dir, filename)
                target_file = os.path.join(target_word_dir, filename)
                shutil.copy2(source_file, target_file)
                print(f"Copied {filename}")
        
        # Copy all relationship files for headers and footers
        template_rels_dir = os.path.join(template_word_dir, "_rels")
        target_rels_dir = os.path.join(target_word_dir, "_rels")
        
        if os.path.exists(template_rels_dir):
            # Ensure target directory exists
            os.makedirs(target_rels_dir, exist_ok=True)
            
            for filename in os.listdir(template_rels_dir):
                if "header" in filename or "footer" in filename:
                    source_file = os.path.join(template_rels_dir, filename)
                    target_file = os.path.join(target_rels_dir, filename)
                    shutil.copy2(source_file, target_file)
                    print(f"Copied relationship file: {filename}")
        
        # Copy all media files from template to target
        template_media_dir = os.path.join(template_word_dir, "media")
        target_media_dir = os.path.join(target_word_dir, "media")
        
        if os.path.exists(template_media_dir):
            # Ensure target directory exists
            os.makedirs(target_media_dir, exist_ok=True)
            
            for filename in os.listdir(template_media_dir):
                source_file = os.path.join(template_media_dir, filename)
                # Only copy if it's a file (not a directory)
                if os.path.isfile(source_file):
                    target_file = os.path.join(target_media_dir, filename)
                    shutil.copy2(source_file, target_file)
                    print(f"Copied media file: {filename}")
        
        # Copy content types file
        template_types_file = os.path.join(template_dir, "[Content_Types].xml")
        target_types_file = os.path.join(target_dir, "[Content_Types].xml")
        
        if os.path.exists(template_types_file):
            shutil.copy2(template_types_file, target_types_file)
            print("Copied content types file")
        
        # Now update the document.xml by copying sectPr from template
        template_doc_file = os.path.join(template_dir, "word", "document.xml")
        target_doc_file = os.path.join(target_dir, "word", "document.xml")
        
        if os.path.exists(template_doc_file) and os.path.exists(target_doc_file):
            try:
                # Read files as XML
                template_tree = ET.parse(template_doc_file)
                target_tree = ET.parse(target_doc_file)
                
                template_root = template_tree.getroot()
                target_root = target_tree.getroot()
                
                # Get section properties from template
                template_sect_pr = None
                for body in template_root.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body"):
                    for sect_pr in body.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"):
                        template_sect_pr = sect_pr
                
                if template_sect_pr is not None:
                    # Remove any existing sectPr in target
                    for body in target_root.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body"):
                        for sect_pr in body.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"):
                            body.remove(sect_pr)
                    
                    # Import the template sectPr to target
                    template_sect_pr_str = ET.tostring(template_sect_pr, encoding='utf-8').decode('utf-8')
                    
                    # Find the body in target
                    target_body = target_root.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body")
                    if target_body is not None:
                        # Append the template sectPr
                        target_body.append(ET.fromstring(template_sect_pr_str))
                        print("Copied section properties from template")
                    
                    # Save the updated target document
                    target_tree.write(target_doc_file, encoding="utf-8", xml_declaration=True)
                    print("Updated document.xml with template section properties")
            except Exception as e:
                print(f"Error updating document.xml: {e}")
        
        # Create new docx file from the modified directory
        try:
            shutil.make_archive(os.path.splitext(output_path)[0], 'zip', target_dir)
            
            # Rename zip to docx
            zip_path = f"{os.path.splitext(output_path)[0]}.zip"
            if os.path.exists(zip_path):
                # Remove existing output file if it exists
                if os.path.exists(output_path):
                    os.remove(output_path)
                os.rename(zip_path, output_path)
                print(f"Created branded document: {output_path}")
        except Exception as e:
            print(f"Error creating final document: {e}")
            # Try with an alternative approach
            try:
                alt_zip_path = f"{os.path.splitext(output_path)[0]}_alt.zip"
                with ZipFile(alt_zip_path, 'w') as zipf:
                    for root, dirs, files in os.walk(target_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_path, target_dir)
                            zipf.write(file_path, arcname)
                
                if os.path.exists(alt_zip_path):
                    alt_output_path = f"{os.path.splitext(output_path)[0]}_alt.docx"
                    os.rename(alt_zip_path, alt_output_path)
                    print(f"Created alternative branded document: {alt_output_path}")
            except Exception as alt_e:
                print(f"Error with alternative approach: {alt_e}")

def convert_pdf_to_docx(pdf_path, docx_path):
    """Convert PDF file to DOCX format."""
    try:
        # Create a PDF converter object
        cv = Converter(pdf_path)
        # Convert PDF to DOCX with A4 settings
        cv.convert(docx_path, start=0, end=None, pages=None, 
                  layout_kwargs={
                      'margin_top': 72.0,      # 1 inch
                      'margin_bottom': 72.0,   # 1 inch
                      'margin_left': 72.0,     # 1 inch
                      'margin_right': 72.0,    # 1 inch
                      'header_height': 36.0,   # 0.5 inch for header
                      'footer_height': 36.0,   # 0.5 inch for footer
                      'overlap_threshold': 20, # Reduce image overlap warnings
                      'ignore_image_overlap': True
                  })
        cv.close()
        
        try:
            doc = Document(docx_path)
            for section in doc.sections:
                section.different_first_page_header_footer = False
                section.header_distance = Inches(0.5)
                section.footer_distance = Inches(0.5)
                section.page_width = Inches(8.27)  # A4 width
                section.page_height = Inches(11.69)  # A4 height
            doc.save(docx_path)
            print("Optimized PDF conversion settings")
        except Exception as e:
            print(f"Warning: Could not optimize converted document: {e}")
        
        return True
    except Exception as e:
        print(f"Error converting PDF to DOCX: {e}")
        return False

def batch_process():
    """Process all files in input directory."""
    for filename in os.listdir(INPUT_DIR):
        input_path = os.path.join(INPUT_DIR, filename)
        
        # Skip if it's a directory
        if os.path.isdir(input_path):
            continue
            
        # Determine output path (always DOCX)
        base_name = os.path.splitext(filename)[0]
        output_path = os.path.join(OUTPUT_DIR, f"{base_name}.docx")
        
        # Make sure output directory exists
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        
        try:
            if filename.endswith(".docx"):
                print(f"\nProcessing DOCX: {filename}")
                apply_branding_to_docx(input_path, output_path)
                print(f"Completed: {filename}")
                    
            elif filename.endswith(".pdf"):
                print(f"\nProcessing PDF: {filename}")
                # First convert PDF to DOCX
                temp_docx = os.path.join(OUTPUT_DIR, f"{base_name}_converted.docx")
                if convert_pdf_to_docx(input_path, temp_docx):
                    # Now apply branding to the converted DOCX
                    apply_branding_to_docx(temp_docx, output_path)
                    
                    # Clean up temporary DOCX file
                    if os.path.exists(temp_docx):
                        os.remove(temp_docx)
                        
                    print(f"Completed: {filename} (converted to DOCX)")
                else:
                    print(f"Failed to convert PDF: {filename}")
                    # Create fallback copy
                    temp_fallback = os.path.join(OUTPUT_DIR, f"{base_name}_fallback.pdf") 
                    shutil.copy2(input_path, temp_fallback)
                    print(f"Fallback saved as: {temp_fallback}")
            else:
                print(f"\nSkipping unsupported file: {filename}")
                
        except Exception as e:
            print(f"Error processing {filename}: {e}")
            # Create fallback copy
            try:
                ext = os.path.splitext(filename)[1]
                temp_fallback = os.path.join(OUTPUT_DIR, f"{base_name}_fallback{ext}")
                shutil.copy2(input_path, temp_fallback)
                print(f"Fallback saved as: {temp_fallback}")
            except Exception as copy_error:
                print(f"Error creating fallback: {copy_error}")

if __name__ == "__main__":
    batch_process()
    print(f"\nProcessing complete. Files saved to: {OUTPUT_DIR}")